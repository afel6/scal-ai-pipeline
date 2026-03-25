from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io

class PRCWordExporter:
    """
    Macro-Level MS Word (.docx) Generation Engine for the Petroleum Research Center.
    Clones the PRC institutional formatting and inserts Recharts/Matplotlib 
    vector graphs explicitly for executive engineering review.
    """
    def __init__(self, well_name: str, study_type="FINAL REPORT (CCA & SCAL)"):
        self.doc = Document()
        self.well_name = well_name
        self.study_type = study_type
        self._setup_prc_styles()

    def _setup_prc_styles(self):
        """Sets the default font to Times New Roman and configures PRC Header styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

    def generate_title_page(self):
        self.doc.add_paragraph('\n\n\n\n')
        
        prc = self.doc.add_heading('PETROLEUM RESEARCH CENTER (PRC)', level=0)
        prc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph('\n\n')
        title = self.doc.add_heading(f'{self.study_type}\n\nWELL: {self.well_name}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph('\n\n\n\n\n\n')
        footer = self.doc.add_paragraph('Automated Experimental Reservoir Data & Fluid Displacement Simulation\nGenerated via Expert AI Pipeline')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_page_break()

    def add_cca_table(self, data_rows: list):
        """Builds the massive Conventional Core Analysis (CCA) layout."""
        self.doc.add_heading('1. Routine Core Analysis (RCA/CCA) Data', level=2)
        
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sample ID'
        hdr_cells[1].text = 'Depth (ft)'
        hdr_cells[2].text = 'Porosity (frac)'
        hdr_cells[3].text = 'Permeability (mD)'
        hdr_cells[4].text = 'Grain Density (g/cc)'
        
        for row in data_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.get('id', 'N/A'))
            row_cells[1].text = f"{row.get('depth', 0.0):.2f}"
            row_cells[2].text = f"{row.get('porosity', 0.0):.4f}"
            row_cells[3].text = f"{row.get('permeability', 0.0):.2f}"
            row_cells[4].text = f"{row.get('grain_density', 0.0):.3f}"
        
        self.doc.add_paragraph('\n')

    def add_scal_physics_plot(self, sw_array, krw_array, kro_array, sample_depth: float):
        """Draws the SCAL curve using Matplotlib and embeds the high-res PNG directly into MS Word."""
        self.doc.add_heading(f'2. Special Core Analysis (SCAL) - Depth: {sample_depth} ft', level=2)
        
        plt.figure(figsize=(7, 5))
        plt.plot(sw_array, krw_array, label='Krw (Water)', color='blue', linewidth=2)
        plt.plot(sw_array, kro_array, label='Kro (Oil)', color='green', linewidth=2)
        plt.xlabel('Water Saturation (Sw)')
        plt.ylabel('Relative Permeability')
        plt.title('Physics-Informed Neural Network Simulation')
        plt.grid(True, linestyle='--', alpha=0.6)
        plt.legend()
        
        # Save to memory buffer to avoid cluttered disk I/O
        mem_stream = io.BytesIO()
        plt.savefig(mem_stream, format='png', dpi=300)
        mem_stream.seek(0)
        
        self.doc.add_picture(mem_stream, width=Inches(6.0))
        plt.close()
        
    def export(self) -> str:
        filename = f"PRC_Final_Report_{self.well_name}.docx"
        self.doc.save(filename)
        return filename
