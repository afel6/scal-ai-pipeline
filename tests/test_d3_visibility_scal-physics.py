"""D3.1 visibility — batch scal-physics (scal_file_handler.py, physics_sandbox.py,
prc_physics.py, physics_validator.py).

Every test forces the failure (a monkeypatched dependency raises / returns the
bad shape, a degenerate dataset) and asserts what the CALLER sees: the return
value, a raised exception, or a marker in the payload. A log line alone never
passes. No network: the D0 egress guard is armed for the whole run.
"""
import numpy as np
import openpyxl
import pytest
from docx import Document

import physics_sandbox
import physics_validator
import prc_physics
import scal_file_handler as sfh
from physics_sandbox import PhysicsSandbox, PhysicalValidationError, archie_formation_factor
from physics_validator import PhysicsEngineError, PhysicsGuard, PhysicsValidator
from prc_physics import (EndpointProvenanceError, calculate_compressibility_sweep,
                         enrich_json_with_brooks_corey, fit_brooks_corey, provenance_notice)


# ── fixtures ──────────────────────────────────────────────────────────────────

@pytest.fixture
def kr_xlsx(tmp_path):
    """A minimal relative-permeability workbook the KR extractor parses."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sample 1"
    ws.append(["Well", "T1-31"])
    ws.append(["Sw", "Krw", "Kro"])
    for sw, krw, kro in [(0.2, 0.0, 0.9), (0.4, 0.05, 0.5), (0.6, 0.2, 0.2), (0.8, 0.5, 0.0)]:
        ws.append([sw, krw, kro])
    p = tmp_path / "KR_T1-31.xlsx"
    wb.save(p)
    return str(p)


@pytest.fixture
def small_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("Well T1-31 SCAL summary")
    doc.add_table(rows=1, cols=2).rows[0].cells[0].text = "Swi"
    p = tmp_path / "report.docx"
    doc.save(p)
    return str(p)


# ── scal_file_handler ─────────────────────────────────────────────────────────

def test_extractor_imports_are_module_top_not_lazy():
    # A missing/shadowed extractors package now fails at import time (visible),
    # instead of propagating from extract() into extract_file_data's swallow.
    for name in ("MICPExtractor", "KRExtractor", "PCExtractor", "RCALExtractor",
                 "FRFExtractor", "detect_unit", "normalize_value"):
        assert hasattr(sfh, name), name


def test_primary_extractor_crash_is_named_in_the_payload(kr_xlsx, monkeypatch):
    # Force the primary handler to crash the way a broken extractor import did.
    monkeypatch.setattr(sfh.SCALFileHandler, "extract",
                        lambda self: (_ for _ in ()).throw(ImportError("No module named extractors.kr")))
    res = sfh.extract_file_data(kr_xlsx, original_filename="KR_T1-31.xlsx")
    assert res["status"] == "success"                       # the robust parser still read it
    assert res["extractor"] == "robust_fallback"
    assert "ImportError" in res["handler_error"] and "extractors.kr" in res["handler_error"]


def test_primary_extractor_success_is_labelled(kr_xlsx):
    res = sfh.extract_file_data(kr_xlsx, original_filename="KR_T1-31.xlsx")
    assert res["extractor"] == "scal_file_handler"
    assert res["handler_error"] is None if "handler_error" in res else True
    assert "graph_linked" in res and res["graph_error"] is None


def test_ground_truth_inventory_raises_on_unreadable_file(tmp_path):
    bad = tmp_path / "corrupt.xlsx"
    bad.write_bytes(b"PK\x03\x04" + b"\x00" * 64)
    with pytest.raises(RuntimeError, match="corrupt.xlsx"):
        sfh.extract_absolute_file_truth([(str(bad), "corrupt.xlsx")])


def test_units_helper_failure_is_raised_not_written_as_a_file_error(kr_xlsx, monkeypatch):
    monkeypatch.setattr(sfh, "detect_unit",
                        lambda col: (_ for _ in ()).throw(ImportError("No module named hviel.utils.units")))
    with pytest.raises(RuntimeError, match="ImportError"):
        sfh.extract_absolute_file_truth([(kr_xlsx, "KR_T1-31.xlsx")])


def test_citation_cleanup_never_fabricates_a_filename():
    raw = "Porosity 0.21 Source: Company:Well:Sample:Porosity"
    assert sfh.clean_citation_clutter(raw, []) == raw
    assert sfh.clean_citation_clutter(raw, None) == raw
    assert "SCAL_AI_Diagnostic_Test" not in sfh.clean_citation_clutter(raw, [])
    assert "*Source: real.xlsx*" in sfh.clean_citation_clutter(raw, ["real.xlsx"])


def test_ledger_compaction_claims_nothing_it_did_not_check():
    block = ("Source File: \nWorksheet: MICP_TestA\nData Range: Rows 3-20 of column B\n"
             "Extraction Engine: \n")
    out = sfh.compress_traceability_ledger(block, ["Mercury Injection Well T1-31.xls"])
    assert "🔒 Data Integrity Status" in out
    assert "verified against" not in out and "programmatic confidence" not in out
    assert "SCAL_AI_Diagnostic_Test" not in out and "Deterministic Analytical Parser" not in out
    assert "Row 14 Col 2" not in out and "Row 1 Col 4" not in out
    assert "Rows 3-20 of column B" in out                   # the range as cited survives
    assert "Mercury Injection Well T1-31.xls" in out         # source = the uploaded file
    assert "`not stated`" in out                             # engine: the block named none
    out2 = sfh.compress_traceability_ledger(block, None)
    assert "**📄 Source File:** `not stated`" in out2


def test_malformed_protocol2_is_a_violation_in_both_validators():
    inventory = {"sheets_found": ["S1"],
                 "sheet_inventories": [{"sheet_name": "S1", "header_row_raw": "Ka (mD) | Porosity"}]}
    parsed = {"protocol_1_file_open_proof": {"target_sheet": "S1"},
              "protocol_2_header_unit_double_check": "checked, trust me",
              "extracted_data": [{"Ka": 1}]}
    v1 = sfh.validate_extraction_against_inventory(parsed, inventory)
    v2 = sfh.validate_permeability_column_binding(parsed)
    assert v1 and "PROTOCOL_2_MALFORMED" in v1[0]
    assert v2 and "PROTOCOL_2_MALFORMED" in v2[0]
    # Absent protocol_2 stays a clean pass — pinned by test_phase0b_hardening.py.
    assert sfh.validate_permeability_column_binding({}) == []


def test_graph_link_failure_is_in_the_payload(kr_xlsx, monkeypatch):
    import geological_graph
    monkeypatch.setattr(geological_graph, "GeologicalGraph",
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("graph db locked")))
    res = sfh.extract_file_data(kr_xlsx, original_filename="KR_T1-31.xlsx")
    assert res["row_count"] > 0 and res["extractor"] == "scal_file_handler"   # extraction itself succeeded
    assert res["graph_linked"] == 0
    assert "graph db locked" in res["graph_error"]


def test_docx_paragraph_only_fallback_is_marked(small_docx, monkeypatch):
    import file_reader
    monkeypatch.setattr(file_reader, "read_file",
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("table reader broke")))
    res = sfh.extract_file_data(small_docx, original_filename="report.docx")
    assert res["status"] == "success" and "T1-31" in res["extracted"]["raw_text"]
    assert res["extraction_mode"] == "paragraph_only"
    assert any("table reader broke" in w for w in res["warnings"])
    assert any("NOT extracted" in w for w in res["warnings"])


# ── physics_sandbox ───────────────────────────────────────────────────────────

def test_sandbox_config_failure_raises_instead_of_silent_defaults(monkeypatch):
    class _NoSettings:
        pass
    monkeypatch.setattr(physics_sandbox, "settings", _NoSettings())
    with pytest.raises(AttributeError):
        PhysicsSandbox()


def _bad_grid(self, model, params):
    sw = self._Sw.tolist()
    n = len(sw)
    return {"Sw": sw, "Krw": [1.0 - i / n for i in range(n)],      # decreasing: fails
            "Kro": [i / n for i in range(n)]}                     # increasing: fails


def test_brooks_corey_corrected_parameters_describe_the_plotted_curve(monkeypatch):
    monkeypatch.setattr(physics_sandbox.KrCurveFitter, "generate_grid", _bad_grid)
    sw = np.linspace(0.2, 0.8, 10)
    out = PhysicsSandbox().fit_brooks_corey(sw, sw ** 2, (1 - sw) ** 2, swi=0.2, sor=0.2)
    assert out["corrected"] is True
    p = out["parameters"]
    assert p["parameters_source"] == "corrected"
    ep = physics_sandbox.Endpoints(Swi=0.2, Sor=0.2, Krw_max=p["Krw_max"], Kro_max=p["Kro_max"])
    x = np.asarray(out["coordinates"]["x"])
    assert out["coordinates"]["y"][0] == pytest.approx(
        [round(float(v), 6) for v in physics_sandbox.bc_krw(x, ep, p["nw"])], abs=1e-6)
    assert out["coordinates"]["y"][1] == pytest.approx(
        [round(float(v), 6) for v in physics_sandbox.bc_kro(x, ep, p["no"])], abs=1e-6)


def test_archie_health_is_scored_on_the_free_fit_and_only_the_fitted_pair():
    phi = np.linspace(0.05, 0.35, 12)
    out = PhysicsSandbox().fit_archie(phi, archie_formation_factor(phi, a=1.0, m=3.2), "FF")
    assert out["corrected"] is True and 1.3 <= out["parameters"]["m"] <= 2.5
    h = out["health"]
    assert "ARCHIE_M_RANGE" in [v["rule"] for v in h["violations"]]
    assert h["grade"] != "A" and h["score"] < 95
    assert h["rules_checked"] == 2                      # a/m only — no textbook b/n passes
    assert h["bounds_source"]


def test_archie_degenerate_input_refuses_instead_of_textbook_fit():
    with pytest.raises(PhysicalValidationError, match="no parameters were fitted"):
        PhysicsSandbox().fit_archie([0.0, 0.0, 0.3], [0.0, 0.0, 5.0], "FF")
    # Through the dispatch shape app.py uses: the failure is the tool result.
    with pytest.raises(PhysicalValidationError):
        physics_sandbox.run_sandboxed("result = sandbox.fit_archie(x, y, model_type)",
                                      inputs={"sandbox": PhysicsSandbox(), "x": [0.0], "y": [0.0],
                                              "model_type": "RI"})


def test_waxman_smits_clamped_or_failed_fit_is_not_grade_A(monkeypatch):
    from physics_sandbox import waxman_smits_conductivity
    sw = np.linspace(0.2, 1.0, 15)
    consts = dict(cw=5.0, b_coeff=0.045, qv=0.3, f_star=12.0)
    out = PhysicsSandbox().fit_waxman_smits(sw, waxman_smits_conductivity(sw, 3.5, **consts), **consts)
    assert out["corrected"] is True and out["parameters"]["n_star"] == 2.5
    assert "ARCHIE_N_RANGE" in [v["rule"] for v in out["health"]["violations"]]
    assert out["health"]["grade"] != "A"

    monkeypatch.setattr(physics_sandbox, "curve_fit",
                        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("no convergence")))
    out2 = PhysicsSandbox().fit_waxman_smits(sw, waxman_smits_conductivity(sw, 2.0, **consts), **consts)
    assert out2["corrected"] is True
    assert "WAXMAN_SMITS_FIT_FAILED" in [v["rule"] for v in out2["health"]["violations"]]
    assert out2["health"]["grade"] != "A"


def test_corrector_and_validator_share_one_bounds_source(monkeypatch):
    narrow = {"a": (0.5, 1.5), "m": (1.3, 2.5), "b": (0.5, 1.5), "n": (1.8, 2.2)}
    monkeypatch.setattr(PhysicsGuard, "archie_bounds", classmethod(lambda cls, basin="Default": (narrow, "test-basin")))
    sw = np.linspace(0.2, 1.0, 12)
    out = PhysicsSandbox().fit_archie(sw, physics_sandbox.archie_resistivity_index(sw, 1.0, 1.6), "RI")
    assert out["corrected"] is True                       # 1.6 is outside the injected window
    assert 1.8 <= out["parameters"]["n"] <= 2.2           # clamped to THAT window
    assert out["health"]["bounds_source"] == "test-basin"
    assert "ARCHIE_N_RANGE" in [v["rule"] for v in out["health"]["violations"]]


# ── prc_physics ───────────────────────────────────────────────────────────────

def test_cp_failure_reason_is_visible_and_no_lithology_is_deduced():
    rows = [{"Pressure_psi": 0.0, "Porosity_percent": 20.0, "Air_Permeability_md": 50.0},
            {"Pressure_psi": 2000.0, "Porosity_percent": 21.0, "Air_Permeability_md": 50.0}]  # rose
    out = calculate_compressibility_sweep(rows)
    assert out[1]["Pore_Volume_Compressibility_psi_inv"] is None
    assert "cannot exceed initial porosity" in out[1]["Cp_error"]
    assert out[1]["Deduced_Lithology"] == "Unknown Matrix"


def test_cp_physics_guard_is_a_module_top_import_and_its_audit_reaches_the_rows():
    assert prc_physics.PhysicsGuard is physics_validator.PhysicsGuard
    rows = [{"Pressure_psi": 0.0, "Porosity_percent": 20.0},
            {"Pressure_psi": 10.0, "Porosity_percent": 10.0}]     # Cp = 0.05 psi^-1: absurd
    out = calculate_compressibility_sweep(rows)
    assert "CP_CATASTROPHIC" in [v["rule"] for v in out[-1]["_cp_physics_audit"]["violations"]]


def test_unparsable_lab_endpoint_is_warned_and_the_report_path_refuses():
    rows = [{"Water_Saturation_fraction": 0.3, "Krw": 0.05, "Kro": 0.6, "explicit_Swi": "n/a"},
            {"Water_Saturation_fraction": 0.6, "Krw": 0.3, "Kro": 0.1, "explicit_Swi": "n/a"}]
    fit = fit_brooks_corey([dict(r) for r in rows])
    assert fit["parameters"]["Swi"]["source"] == "fitted"
    assert any("explicit_Swi present but unparsable" in w for w in fit["warnings"])
    assert "unparsable" in provenance_notice(fit)
    with pytest.raises(EndpointProvenanceError, match="unparsable"):
        enrich_json_with_brooks_corey([dict(r) for r in rows])


def test_non_physical_measured_swi_is_declared_substituted_not_hidden_in_se():
    rows = [{"Water_Saturation_fraction": 0.3, "Krw": 0.05, "explicit_Swi": 0.95},
            {"Water_Saturation_fraction": 0.9, "Krw": 0.4, "explicit_Swi": 0.95}]
    fit = fit_brooks_corey(rows)
    assert fit["parameters"]["Swi"]["source"] == "substituted"
    assert set(fit["substituted"]) >= {"Swi", "Sor"}
    assert fit["endpoints_used"] == {"Swi": 0.1, "Sor": 0.1}          # the 0.8 denominator, declared


def test_degenerate_regression_is_default_not_fitted():
    rows = [{"Water_Saturation_fraction": 0.5, "Krw": 0.1},
            {"Water_Saturation_fraction": 0.5, "Krw": 0.3}]            # identical Se: no line
    fit = fit_brooks_corey(rows)
    assert fit["parameters"]["nw"]["source"] == "default"
    assert fit["parameters"]["krw_max"]["source"] == "default"
    assert "nw" in fit["defaulted"]


def test_positive_pc_slope_lambda_is_default_and_noted():
    rows = [{"Water_Saturation_fraction": 0.3, "Pc_psi": 1.0},
            {"Water_Saturation_fraction": 0.6, "Pc_psi": 10.0}]        # Pc rises with Se
    fit = fit_brooks_corey(rows)
    assert fit["parameters"]["lambda"] == {"value": 2.0, "source": "default"}
    assert fit["parameters"]["Pd_psi"]["source"] == "fitted"
    assert "lambda" in fit["defaulted"]
    assert any("non-physical" in w for w in fit["warnings"])


def test_clamped_exponent_is_labelled_clamped():
    rows = [{"Water_Saturation_fraction": 0.3, "Krw": 0.5, "Kro": 0.1},
            {"Water_Saturation_fraction": 0.6, "Krw": 0.1, "Kro": 0.5}]  # negative slopes
    fit = fit_brooks_corey(rows)
    assert fit["parameters"]["nw"] == {"value": 0.5, "source": "clamped"}
    assert fit["parameters"]["no"] == {"value": 0.5, "source": "clamped"}
    assert fit["clamped"] == ["krnw_max", "krw_max", "no", "nw"]
    assert "CLAMPED" in provenance_notice(fit)


# ── physics_validator ─────────────────────────────────────────────────────────

@pytest.mark.parametrize("call, rule", [
    (lambda g: g.validate_archie([], [], "RI"), "RI_NO_DATA"),
    (lambda g: g.validate_archie([], [], "FF"), "FF_NO_DATA"),
    (lambda g: g.validate_pc([], []), "PC_NO_DATA"),
    (lambda g: g.validate_compressibility([0.0, 0.0, float("nan")]), "CP_NO_DATA"),
])
def test_empty_input_is_not_grade_A(call, rule):
    h = call(PhysicsGuard()).generate_health_score()
    assert rule in [v["rule"] for v in h["violations"]]
    assert h["grade"] == "F" and h["score"] == 0
    assert "All curves follow" not in h["summary"]


def test_no_rules_evaluated_is_not_grade_A():
    h = PhysicsGuard().generate_health_score()
    assert h["rules_checked"] == 0 and h["grade"] == "N/A"
    assert h["grade"] not in PhysicsSandbox._PASSING_GRADES
    assert "All curves follow" not in h["summary"] and "nothing has been validated" in h["footer"]


def test_archie_bounds_source_is_reported_and_db_failure_is_visible(monkeypatch, caplog, capsys):
    import app
    monkeypatch.setattr(app, "db", lambda *a, **k: (_ for _ in ()).throw(RuntimeError("db down")))
    bounds, source = PhysicsGuard.archie_bounds("Sirte")
    assert bounds == PhysicsGuard.ARCHIE_BOUNDS_DEFAULT
    assert source.startswith("hardcoded (db unavailable")
    with caplog.at_level("WARNING", logger="prc-physics-guard"):
        h = PhysicsGuard().validate_archie_parameters(a=1.0, m=2.0, basin_name="Sirte").generate_health_score()
    assert h["bounds_source"].startswith("hardcoded (db unavailable")
    assert h["rules_checked"] == 2
    assert "db down" in caplog.text and capsys.readouterr().out == ""      # logger, not print


def test_archie_parameters_only_count_what_was_passed():
    h = PhysicsGuard().validate_archie_parameters(n=2.0).generate_health_score()
    assert h["rules_checked"] == 1 and h["grade"] == "A"
    h2 = PhysicsGuard().validate_archie_parameters().generate_health_score()
    assert h2["grade"] == "F" and "ARCHIE_NO_PARAMETERS" in [v["rule"] for v in h2["violations"]]


def test_core_physics_requires_endpoints():
    with pytest.raises(PhysicsEngineError, match="Swi, Sor"):
        PhysicsValidator.validate_core_physics({"Porosity": 0.2})
    assert PhysicsValidator.validate_core_physics({"Swi": 0.2, "Sor": 0.2, "Porosity": 0.2})["Swi"] == 0.2
