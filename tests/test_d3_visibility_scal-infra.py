"""D3.1 visibility — batch scal-infra (report_generator, geological_graph,
rag_database, alerting, llm_insight_generator, config).

Every test forces the failure and asserts what the CALLER sees: a paragraph in
the .docx, a marker key in the hybrid_search dict, a raised exception, the
SMTP calls that must NOT happen, or a repo-anchored settings value. A log line
alone never passes. No network: nothing here touches a provider.
"""
import json
import pathlib

import matplotlib.pyplot as plt
import pytest
from docx import Document

import config
from geological_graph import GeologicalGraph
from report_generator import PRCReportEngine

ROOT = pathlib.Path(__file__).resolve().parents[1]


# ── report_generator ──────────────────────────────────────────────────────────

def _doc_text(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def _table_text(doc):
    return "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)


def test_audit_sentence_reflects_missing_audit():
    """No PHYSICS HEALTH AUDIT row: the report must not claim the data passed."""
    doc = Document()
    PRCReportEngine(db_path=":memory:")._build_audit_section(doc, [("model", "hello", "")])
    text = _doc_text(doc)
    assert "has passed the PRC Physics Watchtower gates" not in text
    assert "NOT been verified" in text


def test_audit_sentence_reflects_failed_status():
    doc = Document()
    msgs = [("model", "PHYSICS HEALTH AUDIT: 40% | STATUS: FAIL\n- porosity > 1", "")]
    PRCReportEngine(db_path=":memory:")._build_audit_section(doc, msgs)
    text = _doc_text(doc)
    assert "did NOT pass" in text and "FAIL" in text
    assert "has passed the PRC Physics Watchtower gates" not in text


def test_audit_sentence_pass_is_derived_from_rows():
    doc = Document()
    msgs = [("model", "PHYSICS HEALTH AUDIT: 95% | STATUS: PASS_WITH_WARNINGS\nnone", "")]
    PRCReportEngine(db_path=":memory:")._build_audit_section(doc, msgs)
    assert "has passed the PRC Physics Watchtower gates" in _doc_text(doc)
    assert "PASS_WITH_WARNINGS" in _table_text(doc)


def test_audit_unparseable_line_is_a_visible_row():
    """A malformed audit line used to vanish (bare except: pass) while the
    sentence above the table still said everything passed."""
    doc = Document()
    msgs = [("model", "PHYSICS HEALTH AUDIT: 90%", "")]          # no STATUS: -> IndexError
    PRCReportEngine(db_path=":memory:")._build_audit_section(doc, msgs)
    assert "UNPARSEABLE" in _table_text(doc)
    assert "did NOT pass" in _doc_text(doc)


def test_plot_payload_unparseable_is_noted_in_docx():
    doc = Document()
    msgs = [("model", "__PRC_PLOT__\n{\"title\": \"Broken\", \"curves\": [", "")]
    PRCReportEngine(db_path=":memory:")._build_analysis_section(doc, msgs)
    assert "[Figure omitted" in _doc_text(doc) and "could not be parsed" in _doc_text(doc)


def test_chart_exception_is_noted_under_heading(monkeypatch):
    eng = PRCReportEngine(db_path=":memory:")

    def boom(data):
        raise RuntimeError("matplotlib exploded (simulated)")
    monkeypatch.setattr(eng, "_draw_chart_for_doc", boom)
    doc = Document()
    payload = {"title": "Kr curves", "curves": [{"x": [0.1, 0.2], "y": [0.5, 0.6]}]}
    msgs = [("model", f"__PRC_PLOT__\n{json.dumps(payload)}\nInterpretation text.", "")]
    eng._build_analysis_section(doc, msgs)
    text = _doc_text(doc)
    assert "Kr curves" in text
    assert "[Figure omitted" in text and "matplotlib exploded" in text
    assert "Interpretation text." in text                     # the message's prose survives


def test_chart_without_curve_shape_is_noted():
    doc = Document()
    msgs = [("model", "__PRC_PLOT__\n" + json.dumps({"title": "Empty", "series": []}), "")]
    PRCReportEngine(db_path=":memory:")._build_analysis_section(doc, msgs)
    assert "no recognised curve shape" in _doc_text(doc)


def test_chart_with_every_point_filtered_is_noted():
    """All y <= 0 on a log axis: an empty figure used to be captioned as a result."""
    doc = Document()
    payload = {"title": "Perm", "yAxis": {"label": "Permeability (mD)", "log": True},
               "curves": [{"x": [0.1, 0.2], "y": [0.0, -1.0]}]}
    msgs = [("model", "__PRC_PLOT__\n" + json.dumps(payload), "")]
    PRCReportEngine(db_path=":memory:")._build_analysis_section(doc, msgs)
    assert "every point was filtered out" in _doc_text(doc)


def test_chart_failure_closes_the_figure(monkeypatch):
    plt.close("all")

    def fail(*a, **k):
        raise OSError("disk full (simulated)")
    monkeypatch.setattr(plt, "savefig", fail)
    payload = {"title": "T", "curves": [{"x": [1, 2], "y": [3, 4]}]}
    with pytest.raises(OSError):
        PRCReportEngine(db_path=":memory:")._draw_chart_for_doc(payload)
    assert plt.get_fignums() == []


def test_chart_success_path_unchanged():
    payload = {"title": "T", "curves": [{"x": [1, 2], "y": [3, 4]}]}
    buf = PRCReportEngine(db_path=":memory:")._draw_chart_for_doc(payload)
    assert buf.getbuffer().nbytes > 0


# ── geological_graph.hybrid_search ────────────────────────────────────────────

class _Raising:
    def query_analog_wells(self, **kw):
        raise RuntimeError("chroma store corrupt (simulated)")


class _Stub:
    def __init__(self):
        self.calls = []

    def query_analog_wells(self, current_porosity, current_perm, n_results=3):
        self.calls.append((current_porosity, current_perm))
        return [{"id": "w1", "context": "c", "historical_data": {}}]


@pytest.fixture
def graph():
    g = GeologicalGraph(":memory:", seed=True)
    yield g
    g.close()


def test_hybrid_vector_query_failure_is_marked_in_result(graph):
    res = graph.hybrid_search("Gialo Formation", porous_range=(0.2, 0.24),
                              perm_range=(100.0, 140.0), retriever=_Raising())
    assert res["vector"] == []
    assert "chroma store corrupt" in res["vector_unavailable"]


def test_hybrid_missing_window_is_marked_not_no_matches(graph):
    stub = _Stub()
    res = graph.hybrid_search("Gialo Formation", retriever=stub)          # the RAG-router shape
    assert stub.calls == []
    assert "petrophysical window" in res["vector_unavailable"]
    res = graph.hybrid_search("Gialo Formation", porous_range=(0.2, 0.24), retriever=stub)
    assert stub.calls == []
    assert "petrophysical window" in res["vector_unavailable"]


def test_hybrid_no_retriever_is_marked(graph):
    res = graph.hybrid_search("Gialo Formation", porous_range=(0.2, 0.24), perm_range=(1.0, 2.0))
    assert res["vector"] == [] and "retriever" in res["vector_unavailable"]


def test_hybrid_real_result_carries_no_marker(graph):
    res = graph.hybrid_search("Gialo Formation", porous_range=(0.2, 0.24),
                              perm_range=(100.0, 140.0), retriever=_Stub())
    assert res["vector"][0]["id"] == "w1"
    assert "vector_unavailable" not in res


# ── rag_database.query_analog_wells ───────────────────────────────────────────

def _rag_with_collection(collection):
    from rag_database import RAGDatabase
    db = RAGDatabase.__new__(RAGDatabase)          # no chroma client: only .collection is used
    db.collection = collection
    return db


def test_analog_query_filter_failure_raises_instead_of_unfiltered_results():
    """The metadata-filtered query fails; the old fallback re-queried WITHOUT the
    +/-20%/50% window and returned those hits in the same 'analog well' shape."""
    class Coll:
        def query(self, **kw):
            if "where" in kw:
                raise RuntimeError("where filter rejected (simulated)")
            return {"ids": [["unfiltered_well"]], "documents": [["far away"]], "metadatas": [[{}]]}
    with pytest.raises(RuntimeError, match="where filter rejected"):
        _rag_with_collection(Coll()).query_analog_wells(current_porosity=0.2, current_perm=100.0)


def test_analog_query_total_failure_raises_not_empty_list():
    class Coll:
        def query(self, **kw):
            raise RuntimeError("store unreadable (simulated)")
    with pytest.raises(RuntimeError, match="store unreadable"):
        _rag_with_collection(Coll()).query_analog_wells(current_porosity=0.2, current_perm=100.0)


# ── alerting.send_alert: STARTTLS ─────────────────────────────────────────────

def test_starttls_failure_never_logs_in_or_sends(monkeypatch):
    import alerting

    calls = []

    class FakeSMTP:
        def __init__(self, *a, **k):
            pass

        def ehlo(self):
            calls.append("ehlo")

        def starttls(self):
            calls.append("starttls")
            raise RuntimeError("STARTTLS extension not supported (simulated)")

        def login(self, *a):
            calls.append("login")

        def send_message(self, msg):
            calls.append("send")

        def quit(self):
            calls.append("quit")

    monkeypatch.setattr(alerting.smtplib, "SMTP", FakeSMTP)
    monkeypatch.setattr(alerting.smtplib, "SMTP_SSL", FakeSMTP)
    for k, v in {"ALERT_WEBHOOK_URL": None, "ALERT_SMTP_HOST": "smtp.local", "ALERT_SMTP_PORT": 587,
                 "ALERT_SMTP_USER": "u", "ALERT_SMTP_PASSWORD": "p", "ALERT_EMAIL_TO": "ops@prc.local"}.items():
        monkeypatch.setattr(alerting.settings, k, v)
    alerting.send_alert("subject", "message")
    assert "starttls" in calls
    assert "login" not in calls and "send" not in calls


# ── llm_insight_generator.MasterEngineerNode ──────────────────────────────────

def test_master_engineer_offline_text_carries_no_fabricated_numbers():
    from llm_insight_generator import MasterEngineerNode
    out = MasterEngineerNode(llm_call=None).analyze_scal_data([{"Porosity_percent": 12.0}])
    assert "### Reservoir Report" in out and "### Visualizer Directive" in out
    for fabricated in ("3.5e-6", "15.0 psi", "18.5%", "3500 psi"):
        assert fabricated not in out
    assert "no analysis was performed" in out.lower()


def test_master_engineer_llm_failure_raises_to_the_pipeline():
    from llm_insight_generator import MasterEngineerNode

    def down(prompt, system_instruction=None, temperature=0.2):
        raise RuntimeError("provider unreachable (simulated)")
    with pytest.raises(RuntimeError, match="provider unreachable"):
        MasterEngineerNode(llm_call=down).analyze_scal_data([{"Porosity_percent": 12.0}])


# ── config ────────────────────────────────────────────────────────────────────

def test_env_file_is_repo_anchored_not_cwd(monkeypatch, tmp_path):
    (tmp_path / ".env").write_text("SCAL_MAX_UPLOAD_MB=999\n", encoding="utf-8")
    monkeypatch.delenv("SCAL_MAX_UPLOAD_MB", raising=False)
    monkeypatch.chdir(tmp_path)
    assert config.Settings().SCAL_MAX_UPLOAD_MB != 999          # the CWD's .env is not read
    env_file = config.Settings.model_config.get("env_file")
    assert pathlib.Path(str(env_file)) == ROOT / ".env"
    assert config.ENV_FILE == ROOT / ".env"
    assert config.ENV_FILE_LOADED == config.ENV_FILE.is_file()


def test_empty_store_dir_anchors_to_repo_not_cwd():
    assert config.Settings(DB_DIR="").DB_DIR == str(config.REPO_ROOT)
    assert config.Settings(CHROMA_DIR="").CHROMA_DIR == str(config.REPO_ROOT)
    s = config.Settings(GRAPH_DB_PATH="", DB_DIR="")
    assert s.GRAPH_DB_PATH is None
    assert pathlib.Path(s.graph_db_path) == config.REPO_ROOT / "geological_graph.sqlite"
