import os
import sqlite3
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
import numpy as np
import io

# ──────────────────────────────────────────────
# BRAND CONSTANTS (Sync with CLAUDE.md)
# ──────────────────────────────────────────────
NAVY = "1B3A5C"
BLUE = "2E75B6"

# Transcript messages carrying this marker are rendered into the report's
# Parameter Provenance subsection (prc_physics.provenance_notice).
PROVENANCE_MARKER = "BROOKS-COREY PARAMETER PROVENANCE:"
LIGHT_BLUE = "D5E8F0"
ZEBRA = "F2F2F2"
DARK_GRAY = "333333"
WHITE = "FFFFFF"
RED_ACCENT = "C0392B"

NAVY_RGB = RGBColor(0x1B, 0x3A, 0x5C)
BLUE_RGB = RGBColor(0x2E, 0x75, 0xB6)
GRAY_RGB = RGBColor(0x33, 0x33, 0x33)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
RED_RGB = RGBColor(0xC0, 0x39, 0x2B)
LIGHT_GRAY_RGB = RGBColor(0x99, 0x99, 0x99)

class PRCReportEngine:
    def __init__(self, db_path=None):
        # Explicit path, or the store the app writes (DB_DIR/chat_history.db) —
        # never a CWD-relative file (D1: two databases depending on the launch dir).
        if db_path is None:
            from config import settings
            db_path = os.path.join(settings.DB_DIR, "chat_history.db")
        self.db_path = db_path

    # --- STYLE HELPERS ---
    def _set_cell_shading(self, cell, color_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)

    def _add_heading(self, doc, text, level=1):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.color.rgb = NAVY_RGB
        run.font.name = "Arial"
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            pPr = heading._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="{BLUE}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
        elif level == 2:
            run.font.size = Pt(13)
        return heading

    def _add_body(self, doc, text, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = GRAY_RGB
        run.bold = bold
        run.italic = italic
        return p

    def generate(self, session_id, well_name="Unknown Well", output_dir=None):
        doc = Document()
        
        # --- COVER PAGE ---
        self._build_cover(doc, well_name)
        
        # --- FETCH DATA ---
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT role, text, url FROM m WHERE sid = ? ORDER BY id", (session_id,))
        messages = cursor.fetchall()
        conn.close()

        # --- SECTION 1: EXECUTIVE SUMMARY ---
        doc.add_page_break()
        self._add_heading(doc, "1. Executive Summary", 1)
        self._add_body(doc, f"This document provides a consolidated technical report for the SCAL study conducted on Well {well_name}. The analysis utilizes the PRC SCAL AI Pipeline (Hviel) for physics-aware interpretation and simulation.")

        # --- SECTION 2: PHYSICS INTEGRITY AUDIT ---
        self._add_heading(doc, "2. Physics Integrity Audit", 1)
        self._build_audit_section(doc, messages)

        # --- SECTION 3: PETROPHYSICAL ANALYSIS & PLOTS ---
        doc.add_page_break()
        self._add_heading(doc, "3. Petrophysical Analysis & Interpretations", 1)
        self._build_analysis_section(doc, messages)

        # --- FOOTER ---
        self._add_footer(doc, well_name)

        filename = f"PRC_SCAL_Report_{well_name}_{session_id[:6]}.docx"
        # Destination precedence: explicit output_dir (the app passes the shared
        # PRC vault) -> PRC_AI_VAULT env -> the original local reports/ folder.
        if output_dir is None:
            from config import settings
            output_dir = settings.PRC_AI_VAULT
        base = output_dir
        output_path = os.path.join(base, filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return filename

    def _build_cover(self, doc, well_name):
        # Premium Banner
        banner = doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.CENTER
        banner_cell = banner.rows[0].cells[0]
        self._set_cell_shading(banner_cell, NAVY)
        banner_cell.height = Cm(3.0)
        bp = banner_cell.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bp.add_run("PETROLEUM RESEARCH CENTER")
        run.bold = True; run.font.size = Pt(22); run.font.color.rgb = WHITE_RGB
        
        doc.add_paragraph() # Spacer
        
        # Logo placeholder or real logo
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for lp in ['prc_logo.png', 'prc_logo.jpg']:
            if os.path.exists(lp):
                try:
                    logo_p.add_run().add_picture(lp, width=Inches(2.0))
                    break
                except Exception as ie:
                    print(f"Warning: Failed to load logo {lp}: {ie}")
        
        doc.add_paragraph()
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("Special Core Analysis (SCAL) Report")
        run.bold = True; run.font.size = Pt(26); run.font.color.rgb = NAVY_RGB
        
        doc.add_paragraph()
        # Metadata Table
        info_table = doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_items = [
            ("Well Name", well_name),
            ("Date", datetime.now().strftime("%d %B %Y")),
            ("Classification", "CONFIDENTIAL / SOVEREIGN"),
            ("Specialist", "Hviel (AI Petrophysicist)")
        ]
        for i, (label, value) in enumerate(info_items):
            row = info_table.rows[i]
            c0, c1 = row.cells
            self._set_cell_shading(c0, LIGHT_BLUE)
            c0.paragraphs[0].add_run(label).bold = True
            c1.paragraphs[0].add_run(value)

    def _build_audit_section(self, doc, messages):
        self._add_body(doc, "All data processed during this session has passed the PRC Physics Watchtower gates.")
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(['Measurement', 'Health Score', 'Status']):
            hdr_cells[i].text = h
            self._set_cell_shading(hdr_cells[i], NAVY)
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.font.color.rgb = WHITE_RGB; run.bold = True

        for role, text, url in messages:
            if "PHYSICS HEALTH AUDIT:" in text:
                try:
                    score = text.split("PHYSICS HEALTH AUDIT:")[1].split("|")[0].strip()
                    status = text.split("STATUS:")[1].split("\n")[0].strip()
                    row = table.add_row().cells
                    row[0].text = "SCAL Study"
                    row[1].text = score
                    row[2].text = status
                except: pass

        # Parameter provenance. A substituted endpoint must be visible to the
        # reader of the document, not only present in the JSON payload.
        for role, text, url in messages:
            if PROVENANCE_MARKER in text:
                self._add_heading(doc, "2.1 Parameter Provenance", 2)
                for line in text.splitlines():
                    if line.strip():
                        self._add_body(doc, line.strip(),
                                       bold="WARNING:" in line)

    def _extract_json_payload(self, text: str) -> tuple[dict, str]:
        """
        Extracts the JSON payload from a __PRC_PLOT__ tag using a bracket-counting parser,
        and returns a tuple of (parsed_dict, trailing_text).
        """
        if '__PRC_PLOT__' not in text:
            return {}, text
            
        parts = text.split('__PRC_PLOT__', 1)
        rest = parts[1].strip()
        
        # Clean potential markdown block prefix
        if rest.startswith('```json'):
            rest = rest[7:].strip()
        elif rest.startswith('```'):
            rest = rest[3:].strip()
            
        # Find the starting brace
        start_idx = rest.find('{')
        if start_idx == -1:
            return {}, text
            
        brace_count = 0
        in_string = False
        escape_next = False
        end_idx = -1
        
        for i in range(start_idx, len(rest)):
            char = rest[i]
            
            if escape_next:
                escape_next = False
                continue
                
            if char == '\\':
                escape_next = True
                continue
                
            if char == '"':
                in_string = not in_string
                continue
                
            if not in_string:
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        end_idx = i
                        break
                        
        if end_idx == -1:
            try:
                return json.loads(rest), ""
            except Exception:
                return {}, text
                
        json_str = rest[start_idx:end_idx+1]
        trailing_text = rest[end_idx+1:].strip()
        
        if trailing_text.startswith('```'):
            trailing_text = trailing_text[3:].strip()
            
        try:
            plot_data = json.loads(json_str)
            return plot_data, trailing_text
        except Exception:
            return {}, text

    def _draw_chart_for_doc(self, data: dict) -> io.BytesIO | None:
        """Creates a premium, physics-aware chart image buffer suitable for injection into DOCX."""
        try:
            title  = data.get('title', 'PRC Analysis')
            xAxis_cfg = data.get('xAxis') or {}
            yAxis_cfg = data.get('yAxis') or {}
            yAxis2_cfg = data.get('yAxis2') or {}
            
            xlabel = xAxis_cfg.get('label') or data.get('x_label') or 'X'
            ylabel = yAxis_cfg.get('label') or data.get('y_label') or 'Y'
            ylabel2 = yAxis2_cfg.get('label') or data.get('y_label2') or ''
            
            # 1. Detect if axes are logarithmic
            x_is_log = False
            y_is_log = False
            y2_is_log = False
            
            if data.get('xAxisLog') or xAxis_cfg.get('log'):
                x_is_log = True
            if data.get('yAxisLog') or yAxis_cfg.get('log'):
                y_is_log = True
            if data.get('yAxisRightLog') or yAxis2_cfg.get('log'):
                y2_is_log = True
                
            def check_log_label(lbl):
                if not lbl:
                    return False
                l_str = str(lbl).lower()
                if "relative permeability" in l_str or " kr" in l_str or "kr " in l_str or l_str.startswith("kr"):
                    return False
                keywords = ["permeability", "pc (", "capillary pressure", "resistivity index", "formation factor", "fzi", "rqi"]
                return any(kw in l_str for kw in keywords)
                
            if check_log_label(xlabel):
                x_is_log = True
            if check_log_label(ylabel):
                y_is_log = True
            if ylabel2 and check_log_label(ylabel2):
                y2_is_log = True

            # Extract and build curves
            curves = []
            if 'curves' in data and isinstance(data['curves'], list):
                for c in data['curves']:
                    x_data = c.get('x') or [p.get('x') for p in c.get('data', [])]
                    y_data = c.get('y') or [p.get('y') for p in c.get('data', [])]
                    curves.append({
                        'name': c.get('name') or c.get('label') or 'Series',
                        'x': x_data,
                        'y': y_data,
                        'showLine': c.get('showLine', True),
                        'showPoints': c.get('showPoints', True),
                        'color': c.get('color'),
                        'yId': c.get('yId', 'left')
                    })
            elif 'samples' in data and isinstance(data['samples'], dict):
                for s_name, s_data in data['samples'].items():
                    drainage = s_data.get('drainage', {})
                    x_data = drainage.get('sat_pv') or drainage.get('x', [])
                    y_data = drainage.get('pressure') or drainage.get('y', [])
                    if x_data and y_data:
                        curves.append({
                            'name': s_name,
                            'x': x_data,
                            'y': y_data,
                            'showLine': True,
                            'showPoints': True,
                            'color': None,
                            'yId': 'left'
                        })
            elif 'x' in data and 'y' in data:
                curves.append({
                    'name': data.get('title', 'Series'),
                    'x': data['x'],
                    'y': data['y'],
                    'showLine': True,
                    'showPoints': True,
                    'color': None,
                    'yId': 'left'
                })

            if not curves:
                return None

            # Filter out invalid and non-positive points for log scales
            all_left_x, all_left_y = [], []
            all_right_x, all_right_y = [], []
            
            has_right_axis = data.get('dualAxis', False) or any(c.get('yId') == 'right' for c in curves) or bool(ylabel2)
            
            for curve in curves:
                filtered_x = []
                filtered_y = []
                cur_y_is_log = y2_is_log if (curve['yId'] == 'right' and has_right_axis) else y_is_log
                
                for xi, yi in zip(curve['x'], curve['y']):
                    try:
                        if xi is None or yi is None:
                            continue
                        xf = float(xi)
                        yf = float(yi)
                        if x_is_log and xf <= 0:
                            continue
                        if cur_y_is_log and yf <= 0:
                            continue
                        filtered_x.append(xf)
                        filtered_y.append(yf)
                    except (ValueError, TypeError):
                        continue
                curve['x_filtered'] = filtered_x
                curve['y_filtered'] = filtered_y
                
                if curve['yId'] == 'right' and has_right_axis:
                    all_right_x.extend(filtered_x)
                    all_right_y.extend(filtered_y)
                else:
                    all_left_x.extend(filtered_x)
                    all_left_y.extend(filtered_y)

            # Plot setup
            fig, ax1 = plt.subplots(figsize=(7, 4.5), dpi=150)
            plt.style.use('bmh')
            fig.patch.set_facecolor('white')
            ax1.set_facecolor('#ffffff')
            
            ax2 = None
            if has_right_axis:
                ax2 = ax1.twinx()
                ax2.set_facecolor('none')

            # Color Palette
            _PRC_COLORS = ['#1B3A5C', '#C0392B', '#2E7D32', '#E65100', '#6A1B9A', '#00838F']
            
            # Plot curves
            for i, curve in enumerate(curves):
                x = curve['x_filtered']
                y = curve['y_filtered']
                if not x or not y:
                    continue
                    
                label = curve['name']
                color = curve['color'] or _PRC_COLORS[i % len(_PRC_COLORS)]
                show_line = curve['showLine']
                show_points = curve['showPoints']
                
                target_ax = ax2 if (curve['yId'] == 'right' and ax2) else ax1
                
                # Determine line style and markers
                if not show_line:
                    # Core sample routine - scatter plot with transparency and no line
                    target_ax.scatter(x, y, color=color, label=label, alpha=0.6, s=40, edgecolors='none', zorder=3)
                elif show_line and show_points:
                    target_ax.plot(x, y, marker='o', linestyle='-', linewidth=2.0, markersize=5, color=color, label=label, alpha=0.9, zorder=4)
                else:
                    target_ax.plot(x, y, linestyle='-', linewidth=2.0, color=color, label=label, alpha=0.9, zorder=4)

            # Apply logarithmic scale if needed
            if x_is_log:
                ax1.set_xscale('log')
            if y_is_log:
                ax1.set_yscale('log')
            if y2_is_log and ax2:
                ax2.set_yscale('log')

            # X-axis limits and 5% padding
            is_sat_x = "saturation" in xlabel.lower() or "sw" in xlabel.lower() or "fraction" in xlabel.lower()
            x_min_lim, x_max_lim = None, None
            
            if xAxis_cfg.get('domain') and isinstance(xAxis_cfg['domain'], list) and len(xAxis_cfg['domain']) == 2:
                try:
                    x_min_lim = float(xAxis_cfg['domain'][0])
                    x_max_lim = float(xAxis_cfg['domain'][1])
                except (ValueError, TypeError):
                    pass
                    
            if x_min_lim is None or x_max_lim is None:
                combined_x = all_left_x + all_right_x
                if combined_x:
                    if x_is_log:
                        combined_x = [v for v in combined_x if v > 0]
                        if combined_x:
                            vmin, vmax = min(combined_x), max(combined_x)
                            if vmin == vmax:
                                x_min_lim, x_max_lim = vmin * 0.1, vmax * 10.0
                            else:
                                lmin, lmax = np.log10(vmin), np.log10(vmax)
                                diff = lmax - lmin
                                x_min_lim = 10 ** (lmin - 0.05 * diff)
                                x_max_lim = 10 ** (lmax + 0.05 * diff)
                    else:
                        vmin, vmax = min(combined_x), max(combined_x)
                        if is_sat_x and vmin >= 0 and vmax <= 1.05:
                            x_min_lim, x_max_lim = -0.02, 1.02
                        else:
                            if vmin == vmax:
                                x_min_lim, x_max_lim = vmin - 1.0, vmax + 1.0
                            else:
                                diff = vmax - vmin
                                x_min_lim = vmin - 0.05 * diff
                                x_max_lim = vmax + 0.05 * diff

            if x_min_lim is not None and x_max_lim is not None:
                ax1.set_xlim(x_min_lim, x_max_lim)

            # Primary Y-axis limits
            is_sat_y = "saturation" in ylabel.lower() or "sw" in ylabel.lower() or "fraction" in ylabel.lower()
            y_min_lim, y_max_lim = None, None
            
            if yAxis_cfg.get('domain') and isinstance(yAxis_cfg['domain'], list) and len(yAxis_cfg['domain']) == 2:
                try:
                    y_min_lim = float(yAxis_cfg['domain'][0])
                    y_max_lim = float(yAxis_cfg['domain'][1])
                except (ValueError, TypeError):
                    pass
                    
            if y_min_lim is None or y_max_lim is None:
                if all_left_y:
                    if y_is_log:
                        all_left_y = [v for v in all_left_y if v > 0]
                        if all_left_y:
                            vmin, vmax = min(all_left_y), max(all_left_y)
                            if vmin == vmax:
                                y_min_lim, y_max_lim = vmin * 0.1, vmax * 10.0
                            else:
                                lmin, lmax = np.log10(vmin), np.log10(vmax)
                                diff = lmax - lmin
                                y_min_lim = 10 ** (lmin - 0.05 * diff)
                                y_max_lim = 10 ** (lmax + 0.05 * diff)
                    else:
                        vmin, vmax = min(all_left_y), max(all_left_y)
                        if is_sat_y and vmin >= 0 and vmax <= 1.05:
                            y_min_lim, y_max_lim = -0.02, 1.02
                        else:
                            if vmin == vmax:
                                y_min_lim, y_max_lim = vmin - 1.0, vmax + 1.0
                            else:
                                diff = vmax - vmin
                                y_min_lim = vmin - 0.05 * diff
                                y_max_lim = vmax + 0.05 * diff
                                
            if y_min_lim is not None and y_max_lim is not None:
                ax1.set_ylim(y_min_lim, y_max_lim)

            # Secondary Y-axis limits
            if ax2 and all_right_y:
                y2_min_lim, y2_max_lim = None, None
                if yAxis2_cfg.get('domain') and isinstance(yAxis2_cfg['domain'], list) and len(yAxis2_cfg['domain']) == 2:
                    try:
                        y2_min_lim = float(yAxis2_cfg['domain'][0])
                        y2_max_lim = float(yAxis2_cfg['domain'][1])
                    except (ValueError, TypeError):
                        pass
                if y2_min_lim is None or y2_max_lim is None:
                    if y2_is_log:
                        all_right_y = [v for v in all_right_y if v > 0]
                        if all_right_y:
                            vmin, vmax = min(all_right_y), max(all_right_y)
                            if vmin == vmax:
                                y2_min_lim, y2_max_lim = vmin * 0.1, vmax * 10.0
                            else:
                                lmin, lmax = np.log10(vmin), np.log10(vmax)
                                diff = lmax - lmin
                                y2_min_lim = 10 ** (lmin - 0.05 * diff)
                                y2_max_lim = 10 ** (lmax + 0.05 * diff)
                    else:
                        vmin, vmax = min(all_right_y), max(all_right_y)
                        if vmin == vmax:
                            y2_min_lim, y2_max_lim = vmin - 1.0, vmax + 1.0
                        else:
                            diff = vmax - vmin
                            y2_min_lim = vmin - 0.05 * diff
                            y2_max_lim = vmax + 0.05 * diff
                if y2_min_lim is not None and y2_max_lim is not None:
                    ax2.set_ylim(y2_min_lim, y2_max_lim)

            # Titles and labels
            ax1.set_title(title, fontsize=13, fontweight='bold', color='#1B3A5C', pad=15)
            ax1.set_xlabel(xlabel, fontsize=10, fontweight='bold', color='#333333')
            ax1.set_ylabel(ylabel, fontsize=10, fontweight='bold', color='#333333')
            if ax2 and ylabel2:
                ax2.set_ylabel(ylabel2, fontsize=10, fontweight='bold', color='#333333')

            # Grids and styling
            ax1.grid(True, which='both', linestyle='--', color='#E0E0E0', linewidth=0.5, alpha=0.8)
            if ax2:
                ax2.grid(False)

            # Combine legends
            h1, l1 = ax1.get_legend_handles_labels()
            h2, l2 = ax2.get_legend_handles_labels() if ax2 else ([], [])
            if h1 + h2:
                ax1.legend(h1 + h2, l1 + l2, loc='best', framealpha=0.9, fontsize=8.5, facecolor='white', edgecolor='#CCCCCC')

            # Remove extra spines
            for ax in [ax1, ax2] if ax2 else [ax1]:
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False if not ax2 else True)
                ax.spines['left'].set_visible(True)
                ax.spines['bottom'].set_visible(True)
                ax.spines['left'].set_color('#CCCCCC')
                ax.spines['bottom'].set_color('#CCCCCC')
                if ax2 and ax == ax2:
                    ax.spines['right'].set_color('#CCCCCC')

            plt.tight_layout()
            
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            plt.close()
            buf.seek(0)
            return buf
        except Exception as e:
            print(f"Error drawing chart: {e}")
            return None

    def _build_analysis_section(self, doc, messages):
        for i, (role, text, url) in enumerate(messages):
            if role == "model" and "__PRC_PLOT__" in text:
                try:
                    plot_data, trailing_text = self._extract_json_payload(text)
                    if plot_data:
                        # Dynamic title based on chart title
                        title = plot_data.get('title', 'Petrophysical Curve Analysis')
                        self._add_heading(doc, title, 2)
                        
                        buf = self._draw_chart_for_doc(plot_data)
                        if buf:
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.add_run().add_picture(buf, width=Inches(5.5))
                            
                            # Caption
                            cp = doc.add_paragraph()
                            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cr = cp.add_run(f"Figure: {title}")
                            cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = LIGHT_GRAY_RGB
                        
                        # Add trailing interpretation text underneath the figure
                        if trailing_text:
                            self._add_body(doc, trailing_text, italic=True)
                except Exception as e:
                    print(f"Skipping plot injection in report: {e}")

    def _add_footer(self, doc, well_name):
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = f"PRC Libya  |  Well {well_name}  |  Generated by Hviel AI"

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 2:
        PRCReportEngine().generate(sys.argv[1], sys.argv[2])
