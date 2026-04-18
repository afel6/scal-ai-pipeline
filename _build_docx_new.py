"""
_build_docx_new.py
==================
PRC-branded Word document builder (standalone helper module).
This is used by document_engines.py as the advanced DOCX generation method.

Provides:
    - build_docx(well, engineer, raw_content) -> str (filename)

Requires: python-docx, matplotlib
"""
import os
import re
import json
import time
from datetime import datetime

from docx import Document
from docx.shared import Inches as DocxInches, Pt as DocxPt, Cm, RGBColor as DocxRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

# ── PRC Brand Constants ──
_NAVY_HEX   = "1B3A5C"
_BLUE_HEX   = "2E75B6"
_LTBLUE_HEX = "D5E8F0"
_ZEBRA_HEX  = "F2F2F2"

_NAVY_RGB   = DocxRGB(0x1B, 0x3A, 0x5C)
_BLUE_RGB   = DocxRGB(0x2E, 0x75, 0xB6)
_GRAY_RGB   = DocxRGB(0x33, 0x33, 0x33)
_WHITE_RGB  = DocxRGB(0xFF, 0xFF, 0xFF)
_LTGRAY_RGB = DocxRGB(0x99, 0x99, 0x99)
_RED_RGB    = DocxRGB(0xC0, 0x39, 0x2B)

_PRC_COLORS = ['#1e3a8a', '#E31E24', '#F59E0B', '#16a34a', '#7c3aed', '#0891b2']


# ── Internal helpers ──

def _shd(cell, hex_color: str):
    """Apply background shade to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tcPr.append(shd)


def _brd_custom(cell, **sides):
    """Apply custom borders to a cell via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for side, attrs in sides.items():
        el = parse_xml(
            f'<w:{side} {nsdecls("w")} w:val="single" '
            f'w:sz=\'{attrs.get("sz", "4")}\' '
            f'w:color=\'{attrs.get("color", "CCCCCC")}\'/>'
        )
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _brd_none(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/><w:left w:val="none"/>'
        f'<w:bottom w:val="none"/><w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _page_num_field(paragraph):
    """Insert an auto page-number XML field into a paragraph."""
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2  = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = 'Arial'
    run.font.size = DocxPt(8)
    run.font.color.rgb = _LTGRAY_RGB


def _heading(doc, text: str, level: int = 1):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = DocxPt(20 if level == 1 else 14)
    h.paragraph_format.space_after  = DocxPt(6)
    r = h.add_run(text)
    r.bold = True
    r.font.name = 'Arial'
    r.font.color.rgb = _NAVY_RGB
    r.font.size = DocxPt(16 if level == 1 else 13 if level == 2 else 11)
    if level == 1:
        h._element.get_or_add_pPr().append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="{_BLUE_HEX}"/>'
            f'</w:pBdr>'
        ))


def _body(doc, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after   = DocxPt(8)
    p.paragraph_format.line_spacing  = 1.15
    r = p.add_run(str(text))
    r.bold   = bold
    r.italic = italic
    r.font.name  = 'Arial'
    r.font.size  = DocxPt(10.5)
    r.font.color.rgb = _GRAY_RGB


def _bullet(doc, text: str):
    bp = doc.add_paragraph(style='List Bullet')
    bp.clear()
    r = bp.add_run(str(text))
    r.font.name  = 'Arial'
    r.font.size  = DocxPt(10)
    r.font.color.rgb = _GRAY_RGB


def _styled_table(doc, headers: list, rows: list, col_widths_inches: list = None):
    """Render a PRC-branded data table."""
    ncols = len(headers)
    if not ncols:
        return
    col_w = col_widths_inches or [round(6.0 / ncols, 2)] * ncols

    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = DocxInches(col_w[i] if i < len(col_w) else col_w[-1])

    # Header row
    hdr_row = table.rows[0]
    hdr_row.height = Cm(0.9)
    for ci, h in enumerate(headers):
        cell = hdr_row.cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shd(cell, _NAVY_HEX)
        _brd_custom(cell, top={'sz': '2'}, bottom={'sz': '2'},
                    left={'sz': '2'}, right={'sz': '2'})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        r.bold = True
        r.font.size  = DocxPt(9)
        r.font.color.rgb = _WHITE_RGB
        r.font.name  = 'Arial'

    # Data rows
    for ri, row_data in enumerate(rows):
        row_obj = table.rows[ri + 1]
        row_obj.height = Cm(0.72)
        for ci, val in enumerate(row_data):
            cell = row_obj.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _brd_custom(cell, top={'sz': '2'}, bottom={'sz': '2'},
                        left={'sz': '2'}, right={'sz': '2'})
            if ri % 2 == 1:
                _shd(cell, _ZEBRA_HEX)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size  = DocxPt(9)
            r.font.color.rgb = _GRAY_RGB
            r.font.name  = 'Arial'


def _draw_chart_for_doc(data: dict):
    """
    Renders a PRC-branded chart from plot dict and returns an in-memory PNG buffer.
    Compatible with the __PRC_PLOT__ JSON schema used throughout the pipeline.
    """
    try:
        fig, ax = plt.subplots(figsize=(7, 4.5), dpi=150)
        plt.style.use('bmh')
        fig.patch.set_facecolor('white')
        ax.set_facecolor('#fdfdfd')

        title  = data.get('title',   'PRC Analysis')
        xlabel = data.get('x_label', 'X')
        ylabel = data.get('y_label', 'Y')

        curves = data.get('curves')
        if curves and isinstance(curves, list):
            for i, curve in enumerate(curves):
                cx, cy = curve.get('x', []), curve.get('y', [])
                n = min(len(cx), len(cy))
                ax.plot(cx[:n], cy[:n], marker='o', linestyle='-', linewidth=2.5,
                        markersize=7, color=_PRC_COLORS[i % len(_PRC_COLORS)],
                        label=curve.get('label', f'Series {i + 1}'))
            ax.legend(fontsize=9, framealpha=0.85)
        else:
            x, y = data.get('x', []), data.get('y', [])
            n = min(len(x), len(y))
            ax.plot(x[:n], y[:n], marker='o', linestyle='-',
                    color='#1e3a8a', linewidth=2.5, markersize=8)

        ax.set_title(title,  fontsize=14, fontweight='bold', color='#1B3A5C')
        ax.set_xlabel(xlabel, fontsize=11, fontweight='bold')
        ax.set_ylabel(ylabel, fontsize=11, fontweight='bold')
        ax.grid(True, linestyle='--', alpha=0.4)

        buf = io.BytesIO()
        fig.savefig(buf, format='png', bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception as e:
        print(f"[_build_docx_new] Chart error: {e}")
        return None


# ── Public entry point ──

def build_docx(well: str, engineer: str, raw_content: str) -> str:
    """
    Build a PRC-branded Word document from raw Claude/Gemini JSON content.

    Args:
        well:        Well or project name used in the header.
        engineer:    Authorizing engineer name for the cover page.
        raw_content: Raw text from the LLM (may include ```json fences).

    Returns:
        Filename of the saved .docx file.
    """
    # ── Parse JSON payload ──
    text = raw_content.strip().replace('```json', '').replace('```', '').strip()
    try:
        data = json.loads(text)
    except (json.JSONDecodeError, ValueError):
        m = re.search(r'\{.*\}', text, re.DOTALL)
        try:
            data = json.loads(m.group(0)) if m else None
        except Exception:
            data = None

    doc = Document()

    # ── Global styles ──
    ns = doc.styles['Normal']
    ns.font.name = 'Arial'
    ns.font.size = DocxPt(10.5)
    ns.font.color.rgb = _GRAY_RGB
    ns.paragraph_format.space_after = DocxPt(6)
    for lvl in range(1, 4):
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = 'Arial'
        hs.font.color.rgb = _NAVY_RGB
    for sec in doc.sections:
        sec.top_margin    = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin   = Cm(2.54)
        sec.right_margin  = Cm(2.54)

    # ══════════════════ COVER PAGE ══════════════════

    # Navy top banner
    banner = doc.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    bc = banner.rows[0].cells[0]
    _shd(bc, _NAVY_HEX)
    bc.height = Cm(3.0)
    bp = bc.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.space_before = DocxPt(18)
    run = bp.add_run('PETROLEUM RESEARCH CENTER')
    run.bold = True
    run.font.size = DocxPt(22)
    run.font.color.rgb = _WHITE_RGB
    run.font.name = 'Arial'
    bp2 = bc.add_paragraph()
    bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = bp2.add_run('Tripoli, Libya')
    run2.font.size = DocxPt(12)
    run2.font.color.rgb = DocxRGB(0xAA, 0xCC, 0xDD)
    run2.font.name = 'Arial'

    doc.add_paragraph()

    # Logo
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for lp in ['prc_logo.png', 'prc_logo.jpg']:
        if os.path.exists(lp):
            try:
                logo_p.add_run().add_picture(lp, width=DocxInches(2.0))
                break
            except Exception:
                pass

    # Blue divider
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p._element.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="8" w:color="{_BLUE_HEX}"/>'
        f'</w:pBdr>'
    ))
    doc.add_paragraph()

    # Title
    doc_title = data.get('title', f'PRC Report — {well}') if data else f'PRC Report — {well}'
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr2 = title_p.add_run(doc_title)
    tr2.bold = True
    tr2.font.size = DocxPt(24)
    tr2.font.color.rgb = _NAVY_RGB
    tr2.font.name = 'Arial'

    # Info table
    doc.add_paragraph()
    report_num = (data.get('report_number', f'PRC-SCAL-{datetime.now().year}-AUTO')
                  if data else f'PRC-SCAL-{datetime.now().year}-AUTO')
    info_items = [
        ('Well / Project',    well.upper()),
        ('Prepared By',       engineer),
        ('Report No.',        report_num),
        ('Date',              datetime.now().strftime('%d %B %Y')),
        ('Classification',    'CONFIDENTIAL \u2014 PRC Internal Use Only'),
    ]
    itbl = doc.add_table(rows=len(info_items), cols=2)
    itbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    itbl.autofit   = False
    for row in itbl.rows:
        row.cells[0].width = DocxInches(2.2)
        row.cells[1].width = DocxInches(4.0)
    for i, (label, value) in enumerate(info_items):
        row = itbl.rows[i]
        row.height = Cm(0.65)
        c0 = row.cells[0]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shd(c0, _LTBLUE_HEX)
        p0 = c0.paragraphs[0]
        p0.space_before = DocxPt(2)
        p0.space_after  = DocxPt(2)
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.name = 'Arial'
        r0.font.size = DocxPt(10)
        r0.font.color.rgb = _NAVY_RGB
        c1 = row.cells[1]
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        p1.space_before = DocxPt(2)
        p1.space_after  = DocxPt(2)
        r1 = p1.add_run(str(value))
        r1.font.name = 'Arial'
        r1.font.size = DocxPt(10)
        r1.font.color.rgb = _RED_RGB if i == 4 else _GRAY_RGB
        for cell in [c0, c1]:
            _brd_custom(cell, top={'sz': '2', 'color': 'CCCCCC'},
                        bottom={'sz': '2', 'color': 'CCCCCC'},
                        left={'sz': '2',  'color': 'CCCCCC'},
                        right={'sz': '2', 'color': 'CCCCCC'})

    doc.add_page_break()

    # ══════════════════ HEADER & FOOTER ══════════════════
    section = doc.sections[-1]
    section.different_first_page_header_footer = False

    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    fp._element.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{_BLUE_HEX}"/>'
        f'</w:pBdr>'
    ))
    r1a = fp.add_run('Petroleum Research Center \u2014 Confidential')
    r1a.font.name = 'Arial'
    r1a.font.size = DocxPt(8)
    r1a.font.color.rgb = _LTGRAY_RGB
    r2a = fp.add_run('    |    Page ')
    r2a.font.name = 'Arial'
    r2a.font.size = DocxPt(8)
    r2a.font.color.rgb = _LTGRAY_RGB
    _page_num_field(fp)

    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.clear()
    hp._element.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{_BLUE_HEX}"/>'
        f'</w:pBdr>'
    ))
    hr2 = hp.add_run(f'PRC SCAL  |  Well {well}  |  {report_num}')
    hr2.font.name = 'Arial'
    hr2.font.size = DocxPt(7.5)
    hr2.font.color.rgb = _LTGRAY_RGB
    hr2.italic = True

    # ══════════════════ DOCUMENT BODY ══════════════════
    if data:
        sc = [0]
        for sect in data.get('sections', []):
            lvl  = min(max(int(sect.get('level', 1)), 1), 3)
            htxt = sect.get('heading', '')
            htxt = re.sub(r'^\d+[\.\-\s]+', '', htxt).strip()
            if lvl == 1:
                sc[0] += 1
                htxt = f'{sc[0]}. {htxt}'
            _heading(doc, htxt, level=lvl)

            for para in sect.get('paragraphs', []):
                para_str = str(para).strip()
                if not para_str:
                    continue
                if '__PRC_PLOT__' in para_str:
                    try:
                        pjs = para_str.split('__PRC_PLOT__', 1)[1].strip().replace('```json', '').replace('```', '').strip()
                        chart_data, _ = json.JSONDecoder().raw_decode(pjs)
                        buf = _draw_chart_for_doc(chart_data)
                        if buf:
                            p2 = doc.add_paragraph()
                            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p2.add_run().add_picture(buf, width=DocxInches(6.0))
                            cp = doc.add_paragraph()
                            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cr = cp.add_run(f"Figure: {chart_data.get('title', 'Petrophysical Analysis')}")
                            cr.italic = True
                            cr.font.size = DocxPt(9)
                            cr.font.color.rgb = _GRAY_RGB
                            continue
                    except Exception:
                        pass
                _body(doc, para_str)

            for bul in sect.get('bullets', []):
                _bullet(doc, str(bul))

        for tbl in data.get('tables', []):
            doc.add_paragraph()
            hdrs  = tbl.get('headers', [])
            rows_ = tbl.get('rows',    [])
            if not hdrs:
                continue
            if tbl.get('caption'):
                cp = doc.add_paragraph()
                cr = cp.add_run(tbl['caption'])
                cr.bold = True
                cr.italic = True
                cr.font.name = 'Arial'
                cr.font.size = DocxPt(10)
                cr.font.color.rgb = _NAVY_RGB
            nc = len(hdrs)
            _styled_table(doc, hdrs, rows_, [round(6.0 / nc, 2)] * nc)
            doc.add_paragraph()

    else:
        # Markdown fallback
        _heading(doc, 'ANALYSIS & FINDINGS', level=1)
        for line in raw_content.split('\n'):
            line = line.replace('**', '').strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('#'):
                _heading(doc, line.lstrip('#').strip(),
                         level=min(len(line) - len(line.lstrip('#')), 3))
            elif line.startswith(('* ', '- ')):
                _bullet(doc, line[2:])
            else:
                _body(doc, line)

    # ── Signature block ──
    doc.add_paragraph()
    doc.add_paragraph()
    sep2 = doc.add_paragraph()
    sep2._element.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="8" w:color="{_BLUE_HEX}"/>'
        f'</w:pBdr>'
    ))
    sig = doc.add_table(rows=3, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in sig.rows:
        row.cells[0].width = DocxInches(3.0)
        row.cells[1].width = DocxInches(3.0)
    sig_data = [
        ('Prepared by:', engineer,                     'Hviel \u2014 PRC AI Petrophysical Specialist'),
        ('Reviewed by:', 'PRC Technical Review Board', 'Chief Petrophysicist'),
    ]
    for ci, (role, name, ttl) in enumerate(sig_data):
        for ri, lbl in enumerate([role, name, ttl]):
            p_ = sig.rows[ri].cells[ci].paragraphs[0]
            r_ = p_.add_run(lbl)
            r_.font.name = 'Arial'
            if ri == 0:
                r_.font.size = DocxPt(9)
                r_.font.color.rgb = _LTGRAY_RGB
            elif ri == 1:
                r_.bold = True
                r_.font.size = DocxPt(10)
                r_.font.color.rgb = _NAVY_RGB
            else:
                r_.font.size = DocxPt(9)
                r_.font.color.rgb = _GRAY_RGB
    for row in sig.rows:
        for cell in row.cells:
            _brd_none(cell)

    fname = f'PRC_Report_{int(time.time())}.docx'
    doc.save(fname)
    return fname
